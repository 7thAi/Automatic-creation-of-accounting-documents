"""
Модуль для заполнения приложений Word документами с фотографиями.
"""
import re
from pathlib import Path
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
from io import BytesIO


class PrilozhenieFiller:
    """Заполняет Word-документы приложений фотографиями и подписями."""

    EXTENSIONS = {'.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff'}

    PHOTO_WIDTH = Cm(12.36)
    PHOTO_HEIGHT = Cm(6.49)
    FONT_NAME = "Times New Roman"
    FONT_SIZE = Pt(14)
    CELL_WIDTH = Cm(12)

    HEADERS_MAP = {
        "1. ДТ": "Фотофиксация нарушений, установленных при обследовании содержания и благоустройства дворовых территорий и внутриквартальных проездов ГБУ",
        "2. МКД": "Фотофиксация нарушений, установленных при внешнем обследовании санитарного и технического состояния многоквартирных домов, находящихся в управлении ГБУ",
        "3. ОДХ": "Фотофиксация нарушений, установленных при обследовании содержания объектов дорожного хозяйства ГБУ",
        "4. ОО": "Фотофиксация нарушений, установленных при обследовании содержания и благоустройства объектов озеленения ГБУ"
    }

    VIOLATION_MAP = {
        "1. Проезд АБП": "локальные разрушения на проездах",
        "2. ДТС АБП": "локальные разрушения на ДТС",
        "3. Борт АБП": "локальные разрушения бортового камня",
        "4. ИДН АБП": "неудовлетворительное состояние ИДН",
        "5. Подпорка АБП": "неудовлетворительное состояние подпорной стены",
        "6. Санитарка АБП": "не убрана, не очищена территория",
        "1. Покрытие ДП": "неудовлетворительное состояние покрытия ДП",
        "2. МАФ ДП": "неудовлетворительное состояние МАФ на ДП",
        "3. Табличка ДП": "отсутствует информационная табличка на ДП",
        "4. Санитарка ДП": "не убрана, не очищена территория на ДП",
        "1. Покрытие СП": "неудовлетворительное состояние покрытия СП",
        "2. МАФ СП": "неудовлетворительное состояние МАФ на СП",
        "3. Табличка СП": "отсутствует информационная табличка на СП",
        "4. Санитарка СП": "не убрана, не очищена территория на СП",
        "4. КП": "неудовлетворительное состояние КП",
        "5. МАФ ДТ": "неудовлетворительное состояние МАФ на ДТ",
        "6. Газон": "неудовлетворительное состояние газона и зеленых насаждений",
        "1. Тех": "неудовлетворительное техническое состояние входной группы",
        "2. Сан": "неудовлетворительное санитарное состояние входной группы",
        "2. Отмостка": "неудовлетворительное состояние отмостки",
        "3. Цоколь": "неудовлетворительное состояние цоколя",
        "4. Ливневка": "неудовлетворительное состояние ливневых стоков",
        "5. Надписи": "наклейки, надписи на фасаде здания",
        "6. Сосульки": "наличие снежных масс и сосулек на кровлях и выступающих элементах фасадов зданий",
        "1. Проезд ОДХ": "локальные разрушения на проездах",
        "2. Тротуар ОДХ": "локальные разрушения на тротуарах",
        "3. Борт ОДХ": "локальные разрушения бортового камня",
        "4. ИДН ОДХ": "неудовлетворительное состояние ИДН",
        "5. Санитарка ОДХ": "не убрана, не очищена территория",
        "2. Тех ограждения": "неудовлетворительное техническое состояние ограждений",
        "3. Сан ограждения": "неудовлетворительное санитарное состояние ограждений",
        "4. Сан дорожные знаки": "неудовлетворительное санитарное состояние дорожных знаков",
        "5. МАФ ОДХ": "неудовлетворительное состояние МАФ на ОДХ",
        "1. Проезд ОО": "локальные разрушения на проездах",
        "2. ДТС ОО": "локальные разрушения на ДТС",
        "3. Борт ОО": "локальные разрушения бортового камня",
        "4. ИДН ОО": "неудовлетворительное состояние ИДН",
        "5. Подпорка ОО": "неудовлетворительное состояние подпорной стены",
        "6. Санитарка ОО": "не убрана, не очищена территория",
        "1. Покрытие ОО ДП": "неудовлетворительное состояние покрытия ДП",
        "2. МАФ ОО ДП": "неудовлетворительное состояние МАФ на ДП",
        "3. Табличка ОО ДП": "отсутствует информационная табличка на ДП",
        "4. Санитарка ОО ДП": "не убрана, не очищена территория на ДП",
        "1. Покрытие ОО СП": "неудовлетворительное состояние покрытия СП",
        "2. МАФ ОО СП": "неудовлетворительное состояние МАФ на СП",
        "3. Табличка ОО СП": "отсутствует информационная табличка на СП",
        "4. Санитарка ОО СП": "не убрана, не очищена территория на СП",
        "4. МАФ ОО": "неудовлетворительное состояние МАФ на ОО",
        "5. Газон ОО": "неудовлетворительное состояние газона и зеленых насаждений"
    }

    def _extract_gbu_short_name(self, gbu_name: str) -> str:
        """
        Извлекает короткое название ГБУ из полного названия.
        Например: "ГБУ «Автомобильные дороги ЦАО»" -> "«Автомобильные дороги ЦАО»"
        
        Args:
            gbu_name: Полное название ГБУ.
            
        Returns:
            Короткое название в кавычках.
        """
        match = re.search(r'«(.+?)»', gbu_name)
        if match:
            return f"«{match.group(1)}»"
        return gbu_name

    def _create_title_page(self, doc: Document, gbu_name: str, app_number: int) -> None:
        """
        Создает титульную страницу с номером приложения.
        
        Args:
            doc: Документ Word.
            gbu_name: Название ГБУ.
            app_number: Номер приложения.
        """
        # Удаляем все пустые параграфы в начале документа
        while doc.paragraphs and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)
        
        # Номер приложения справа
        p_number = doc.add_paragraph()
        p_number.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p_number.add_run(f"Приложение № {app_number}")
        run.font.name = self.FONT_NAME
        run.font.size = self.FONT_SIZE
        run.bold = True
        p_number.paragraph_format.line_spacing = 1.5

    def _create_section_header(self, doc: Document, folder: str, gbu_name: str) -> None:
        """
        Создает заголовок для каждой секции (ДТ, МКД, ОДХ, ОО).
        
        Args:
            doc: Документ Word.
            folder: Папка (1. ДТ, 2. МКД, и т.д.).
            gbu_name: Название ГБУ.
        """
        gbu_short = self._extract_gbu_short_name(gbu_name)
        header_template = self.HEADERS_MAP.get(folder, "")
        
        if header_template:
            header_text = header_template + " " + gbu_short
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(header_text)
            run.font.name = self.FONT_NAME
            run.font.size = self.FONT_SIZE
            run.bold = True
            p.paragraph_format.line_spacing = 1.5

    def _create_table_for_photos(self, doc: Document, photo_count: int, left_only: bool = False) -> None:
        """
        Создает таблицу для фотографий с нужным количеством строк.
        
        Args:
            doc: Документ Word.
            photo_count: Количество фотографий.
            left_only: Если True, каждое фото занимает отдельную строку (только левая колонка),
                      если False, два фото в одной строке, подписи во второй строке.
        """
        if left_only:
            # Для приложения устранения: каждое фото + подпись = 2 строки на фото
            rows = photo_count * 2
        else:
            # Для обычного приложения: два фото в одной строке, две подписи во второй
            rows = (photo_count + 1) // 2 * 2  # Округляем вверх для четного количества строк
        
        table = doc.add_table(rows=rows, cols=2)
        table.style = 'Table Grid'
        
        # Устанавливаем интервал строк = 1
        for row in table.rows:
            for cell in row.cells:
                cell.width = self.CELL_WIDTH
                # Устанавливаем интервал в параграфах ячейки
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1

    def fill_prilozhenie(self, template_path: Path, photo_root: Path, save_path: Path,
                         gbu_name: str = None, app_number: int = None,
                         show_progress: bool = False) -> None:
        """
        Заполняет приложение фотографиями (обе колонки).
        
        Args:
            template_path: Путь к шаблону документа (пустой файл).
            photo_root: Корневая папка с фотографиями.
            save_path: Путь для сохранения результата.
            gbu_name: Название ГБУ для вставки в заголовки.
            app_number: Номер приложения.
            show_progress: Показывать прогресс выполнения.
        """
        # Открываем шаблон
        if template_path and template_path.exists():
            doc = Document(str(template_path))
        else:
            doc = Document()
        
        if gbu_name and app_number:
            self._create_title_page(doc, gbu_name, app_number)
        
        folders = ["1. ДТ", "2. МКД", "3. ОДХ", "4. ОО"]
        first_section = True
        
        for folder in folders:
            folder_path = photo_root / folder
            
            if not folder_path.exists():
                if show_progress:
                    print(f"  Папка не существует: {folder_path}")
                continue
            
            photos = self._collect_photos(folder_path)
            
            if not photos:
                if show_progress:
                    print(f"  {folder}: нет фотографий")
                continue
            
            # Добавляем разрыв страницы перед каждой секцией, кроме первой
            if not first_section:
                doc.add_page_break()
            first_section = False
            
            # Добавляем заголовок секции
            if gbu_name:
                self._create_section_header(doc, folder, gbu_name)
            
            # Добавляем таблицу
            self._create_table_for_photos(doc, len(photos), left_only=False)
            table = doc.tables[-1]
            
            # Заполняем таблицу
            self._fill_table(table, photos, left_only=False)
            
            if show_progress:
                print(f"  {folder}: {len(photos)} фото")
        
        doc.save(str(save_path))
        if show_progress:
            print(f"  {save_path.name}: 100% заполнено")

    def fill_prilozhenie_ustraneniya(self, template_path: Path, photo_root: Path, 
                                      save_path: Path, gbu_name: str = None, 
                                      app_number: int = None, show_progress: bool = False) -> None:
        """
        Заполняет приложение устранения (только левая колонка).
        
        Args:
            template_path: Путь к шаблону документа (пустой файл).
            photo_root: Корневая папка с фотографиями.
            save_path: Путь для сохранения результата.
            gbu_name: Название ГБУ для вставки в заголовки.
            app_number: Номер приложения.
            show_progress: Показывать прогресс выполнения.
        """
        # Открываем шаблон
        if template_path and template_path.exists():
            doc = Document(str(template_path))
        else:
            doc = Document()
        
        if gbu_name and app_number:
            self._create_title_page(doc, gbu_name, app_number)
        
        folders = ["1. ДТ", "2. МКД", "3. ОДХ", "4. ОО"]
        first_section = True
        
        for folder in folders:
            folder_path = photo_root / folder
            
            if not folder_path.exists():
                if show_progress:
                    print(f"  Папка не существует: {folder_path}")
                continue
            
            photos = self._collect_photos(folder_path)
            
            if not photos:
                if show_progress:
                    print(f"  {folder}: нет фотографий")
                continue
            
            # Добавляем разрыв страницы перед каждой секцией, кроме первой
            if not first_section:
                doc.add_page_break()
            first_section = False
            
            # Добавляем заголовок секции
            if gbu_name:
                self._create_section_header(doc, folder, gbu_name)
            
            # Добавляем таблицу
            self._create_table_for_photos(doc, len(photos), left_only=True)
            table = doc.tables[-1]
            
            # Заполняем таблицу
            self._fill_table(table, photos, left_only=True)
            
            if show_progress:
                print(f"  {folder}: {len(photos)} фото")
        
        doc.save(str(save_path))
        if show_progress:
            print(f"  {save_path.name}: 100% заполнено")



    def _collect_photos(self, root: Path) -> List[Dict]:
        """
        Собирает информацию о всех фотографиях в папке.
        
        Args:
            root: Корневая папка для сканирования.
            
        Returns:
            Список словарей с информацией о фотографиях.
        """
        result = []
        for item in sorted(root.rglob('*')):
            if item.is_file() and item.suffix.lower() in self.EXTENSIONS:
                subfolder = item.parent.name
                result.append({"path": item, "subfolder": subfolder.strip()})
        return result

    def _fill_table(self, table, photos: List[Dict], left_only: bool = False) -> None:
        """
        Заполняет таблицу фотографиями.
        
        Для left_only=False (обычное приложение):
        - Строка 0: Фото 1 | Фото 2
        - Строка 1: Подпись 1 | Подпись 2
        - Строка 2: Фото 3 | Фото 4
        - Строка 3: Подпись 3 | Подпись 4
        - И так далее...
        
        Для left_only=True (приложение устранения):
        - Строка 0: Фото 1 | (пусто)
        - Строка 1: Подпись 1 | (пусто)
        - Строка 2: Фото 2 | (пусто)
        - Строка 3: Подпись 2 | (пусто)
        - И так далее...
        
        Args:
            table: Таблица Word.
            photos: Список информации о фотографиях.
            left_only: Заполнять только левую колонку.
        """
        if left_only:
            # Приложение устранения: каждое фото занимает 2 строки
            for photo_idx, info in enumerate(photos):
                row_idx = photo_idx * 2
                
                # Вставляем фото в левую колонку
                self._insert_photo(table, row_idx, 0, info["path"])
                
                # Вставляем подпись в левую колонку
                self._insert_caption(table, row_idx + 1, 0, info)
                
                # Очищаем правую колонку
                self._clear_cell(table, row_idx, 1)
                self._clear_cell(table, row_idx + 1, 1)
        else:
            # Обычное приложение: два фото в одной строке, две подписи во второй
            row_idx = 0
            photo_idx = 0
            
            while photo_idx < len(photos):
                # Строка с фотографиями
                left_photo = photos[photo_idx]
                photo_idx += 1
                
                right_photo = None
                if photo_idx < len(photos):
                    right_photo = photos[photo_idx]
                    photo_idx += 1
                
                # Вставляем левое фото
                self._insert_photo(table, row_idx, 0, left_photo["path"])
                
                # Вставляем правое фото или очищаем ячейку
                if right_photo:
                    self._insert_photo(table, row_idx, 1, right_photo["path"])
                else:
                    self._clear_cell(table, row_idx, 1)
                
                # Строка с подписями
                row_idx += 1
                self._insert_caption(table, row_idx, 0, left_photo)
                
                if right_photo:
                    self._insert_caption(table, row_idx, 1, right_photo)
                else:
                    self._clear_cell(table, row_idx, 1)
                
                row_idx += 1

    def _clear_cell(self, table, row: int, col: int) -> None:
        """
        Очищает содержимое ячейки таблицы.
        
        Args:
            table: Таблица Word.
            row: Номер строки.
            col: Номер колонки.
        """
        try:
            cell = table.cell(row, col)
            for p in cell.paragraphs:
                p.clear()
        except Exception:
            pass

    def _insert_photo(self, table, row: int, col: int, path: Path) -> None:
        """
        Вставляет фотографию в ячейку таблицы.
        
        Args:
            table: Таблица Word.
            row: Номер строки.
            col: Номер колонки.
            path: Путь к файлу изображения.
        """
        try:
            with Image.open(path) as img:
                resized_img = self._resize_image(img, 250)
                
                img_bytes = BytesIO()
                resized_img.save(img_bytes, format='PNG')
                img_bytes.seek(0)

                cell = table.cell(row, col)
                # Очищаем все параграфы в ячейке
                for p in cell.paragraphs:
                    p.clear()
                
                # Добавляем фото в первый параграф
                p = cell.paragraphs[0]
                run = p.add_run()
                run.add_picture(img_bytes, width=self.PHOTO_WIDTH, height=self.PHOTO_HEIGHT)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"  Ошибка при вставке фото {path}: {e}")

    def _insert_caption(self, table, row: int, col: int, info: Dict) -> None:
        """
        Вставляет подпись под фотографией в следующий параграф той же ячейки.
        
        Args:
            table: Таблица Word.
            row: Номер строки.
            col: Номер колонки.
            info: Информация о фотографии.
        """
        cell = table.cell(row, col)
        
        # Удаляем все существующие параграфы в ячейке
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)
        
        # Создаем новый параграф для подписи
        p_caption = cell.add_paragraph()
        p_caption.paragraph_format.line_spacing = 1

        address = self._clean_address(info["path"])
        violation = self.VIOLATION_MAP.get(info["subfolder"], "неизвестный тип нарушения").strip()

        run1 = p_caption.add_run("Адрес: ")
        run1.font.name = self.FONT_NAME
        run1.font.size = self.FONT_SIZE
        run1.bold = True

        run2 = p_caption.add_run(address)
        run2.font.name = self.FONT_NAME
        run2.font.size = self.FONT_SIZE
        run2.bold = False

        p_caption.add_run("\n")

        run3 = p_caption.add_run("Нарушение: ")
        run3.font.name = self.FONT_NAME
        run3.font.size = self.FONT_SIZE
        run3.bold = True

        run4 = p_caption.add_run(violation)
        run4.font.name = self.FONT_NAME
        run4.font.size = self.FONT_SIZE
        run4.bold = False

    def _clean_address(self, path: Path) -> str:
        """
        Очищает адрес из имени файла.
        
        Args:
            path: Путь к файлу.
            
        Returns:
            Очищенный адрес.
        """
        name = re.sub(r"\s*\(\d+\)$", "", path.stem)
        return name.replace("_", "/")

    def _resize_image(self, img: Image.Image, target_dpi: int) -> Image.Image:
        """
        Изменяет размер изображения для целевого DPI.
        
        Args:
            img: Исходное изображение.
            target_dpi: Целевое разрешение DPI.
            
        Returns:
            Изменённое изображение.
        """
        original_dpi = img.info.get("dpi", (72, 72))[0]
        x_inch = img.width / original_dpi
        y_inch = img.height / original_dpi
        new_width = int(x_inch * target_dpi)
        new_height = int(y_inch * target_dpi)
        
        # Используем Image.Resampling.LANCZOS для совместимости с Pillow 10+
        try:
            resampling = Image.Resampling.LANCZOS
        except AttributeError:
            # Для старых версий Pillow
            resampling = Image.LANCZOS
            
        resized = img.resize((new_width, new_height), resampling)
        resized.info['dpi'] = (target_dpi, target_dpi)
        return resized
